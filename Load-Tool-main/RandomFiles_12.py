# -*- coding: utf-8 -*-
# Рандомное имя (хеш md5 от имени + timestamp + 'avsw')
# Рандомное содержимое не равное имени (хеш md5 от имени)
# Набор ссылок из скрипта вставляется в каждый файл через пробел

# sudo apt install python3-pip
# pip install pyminizip
# pip3 install python-docx
# pip3 install XlsxWriter
# sudo pip install pathlib
# pip install fpdf
# pip install wheel

# md generator https://dillinger.io/

#  тип      уник    реальн
# ---       ---     ---
# odt       +       +
# rst       +       +
# txt       +       +
# docx      +       +
# docx пар  .       .
# xlsx      +       +
# zip       +       +
# zip пар   +       +
# zip шиф   +       +
# zip мног  .       .
# zip влож  .       .
# md        +       +
# html      +       +
# sh        +       +
# bat       +       +
# pdf 		+		+
# elf 		+ 		+
# deb
# ppt
# exe
# dll

import os
from pathlib import Path
import random
import hashlib
import time
import datetime
import zipfile
# import pyminizip as pyzip
import docx
import xlsxwriter
from fpdf import FPDF
from shutil import copyfile

# from odf.opendocument import OpenDocumentText
# from odf.style import Style, TextProperties
# from odf.text import H, P, Span

import sys

list_all_link = []
list_of_extensions = []
count_pack = 0
nameDirForGenerateFile = ""

staticHash = "578ce80c3b3dfa9ea84d6cd5735ed904 https://www.rbc.ru/society/"
staticLink = "https://www.rbc.ru/society/"


text_rst = '''**reStructuredText**
*Материал из Википедии — свободной энциклопедии*

===================================
Перейти к навигации
===================================
****************
Перейти к поиску
****************
ReStructuredText
================
Изображение логотипа
--------------------
Расширение    .rst

- MIME-тип text/x-rst[1]
- Тип формата текстовые форматы

.. image:: https://3dnews.ru/assets/external/illustrations/2015/09/06/919775/sm.glogo1.600.jpg

Сайт    docutils.sourceforge.io/…​ (англ.)
reStructuredText (сокращение: ReST, расширение файла: .rst) — облегчённый язык разметки. Хорошо применим для создания простых веб-страниц и других документов, а также в качестве вспомогательного языка при написании комментариев в программном коде. 
'''

text_pdf2 = "www.ru"
text_pdf3 = "http://www.yandex"
text_pdf4 = "https://www.rbc.ru/society/"
text_pdf5 = "https://chrome-google.ru"
text_pdf6 = "https://virus-x.3dn.ru/load/virus/3"
text_pdf7 = "https://btsemarkjames.weebly.com/"
text_pdf8 = "http://www.folk-cantabria.com/wp-content/plugins/wp-statistics/includes/classes/gallery_create_page_field.php"
text_pdf9 = "http://www.saipadiesel124.com/wp-content/plugins/imsanity/tmp.php"
text_pdf10 = "https://foundations-admin-1000000787789877014005045757.ml/checkpoint_next.php&quot"
text_pdf11 = "http://cd56640.tmweb.ru/index.php/false/pytn&quot"
text_pdf12 = "http://cd56640.tmweb.ru/index.php/false/pytn"
text_pdf13 = "http://581503.selcdn.ru/sara/rewise.html&quot"
text_pdf14 = "http://grupnew-xxx.duckdns.org/&quot"
text_pdf15 = "http://atp.me/znivmvg2mzfvoduxvzdqmge&quot"
text_pdf16 = "http://pbgnewvvip.com/&quot"
text_pdf17 = "https://pagead2.googlesyndication.com/pagead/js/r20210831/r20110914/abg_lite_fy2019.js"

text_bat = '''@echo off
set vredLink="www.ru"
set vredLink="http://www.yandex"
set vredLink="https://www.rbc.ru/society/"
set vredLink="https://chrome-google.ru"
set vredLink="https://virus-x.3dn.ru/load/virus/3"
set vredLink="https://btsemarkjames.weebly.com/"
set vredLink="http://www.folk-cantabria.com/wp-content/plugins/wp-statistics/includes/classes/gallery_create_page_field.php"
set vredLink="http://www.saipadiesel124.com/wp-content/plugins/imsanity/tmp.php"
set vredLink="https://foundations-admin-1000000787789877014005045757.ml/checkpoint_next.php&quot"
set vredLink="http://cd56640.tmweb.ru/index.php/false/pytn&quot"
set vredLink="http://cd56640.tmweb.ru/index.php/false/pytn"
set vredLink="http://581503.selcdn.ru/sara/rewise.html&quot"
set vredLink="http://grupnew-xxx.duckdns.org/&quot"
set vredLink="http://atp.me/znivmvg2mzfvoduxvzdqmge&quot"
set vredLink="http://pbgnewvvip.com/&quot"
set vredLink="https://pagead2.googlesyndication.com/pagead/js/r20210831/r20110914/abg_lite_fy2019.js"
echo %vredLink%
'''


text_sh = '''#!/bin/bash
PLANET="-"
while [ $PLANET != "EXIT" ]
do
  echo -n "Enter the name of planet: "
  read PLANET
  if [ $PLANET != "EXIT" ]
  then
    echo -n "The $PLANET has "
    case $PLANET in
      Mercury | Venus ) echo -n "no";;
      Earth ) echo -n "one";;
      Mars ) echo -n "two";;
      Jupiter ) echo -n "79";;
      *) echo -n "an unknown number of";;
    esac
  echo " satellite(s)."
  fi
done

'''

text_html = '''<html>
<head>
<meta charset="utf-8"/>
<title>FonColor</title>
</head>
<body>
<script>
function color_fon() {
var abc = "ABCDEF0123456789";
var result = "";
for (var i = 0; i < 6; i++)
result += abc.charAt(Math.floor(Math.random() * abc.length));

document.body.style.background = "#"+result;
document.getElementsByTagName('h2')[0].innerText = "Такой цвет: #" + result;
}
</script>
<h2>123</h2>
<button type="submit" onclick="color_fon()">Сменить цвет фона</button>
<p>VredLink</p>
   <a href="http://www.yandex">v1</a>
   <a href="https://www.rbc.ru/society/">v2</a>
   <a href="https://chrome-google.ru">v3</a>
   <a href="https://virus-x.3dn.ru/load/virus/3">v4</a>
   <a href="https://btsemarkjames.weebly.com/">v5</a>
   <a href="http://www.folk-cantabria.com/wp-content/plugins/wp-statistics/includes/classes/gallery_create_page_field.php">v6</a>
   <a href="http://www.saipadiesel124.com/wp-content/plugins/imsanity/tmp.php">v7</a>
   <a href="https://foundations-admin-1000000787789877014005045757.ml/checkpoint_next.php&quot">v8</a>
   <a href="http://cd56640.tmweb.ru/index.php/false/pytn&quot">v9</a>
   <a href="http://cd56640.tmweb.ru/index.php/false/pytn">v10</a>
   <a href="http://581503.selcdn.ru/sara/rewise.html&quot">v11</a>
   <a href="http://grupnew-xxx.duckdns.org/&quot">v12</a>
   <a href="http://atp.me/znivmvg2mzfvoduxvzdqmge&quot">v13</a>
   <a href="http://pbgnewvvip.com/&quot">v14</a>
   <a href="https://pagead2.googlesyndication.com/pagead/js/r20210831/r20110914/abg_lite_fy2019.js">v15</a>
</body>
</html>
'''

text_md = '''# Dillinger
## _The Last Markdown Editor, Ever_

[![N|Solid](https://cldup.com/dTxpPi9lDf.thumb.png)](https://nodesource.com/products/nsolid)

[![Build Status](https://travis-ci.org/joemccann/dillinger.svg?branch=master)](https://travis-ci.org/joemccann/dillinger)

Dillinger is a cloud-enabled, mobile-ready, offline-storage compatible,
AngularJS-powered HTML5 Markdown editor.

- Type some Markdown on the left
- See HTML in the right
- ✨Magic ✨

## Features

- Import a HTML file and watch it magically convert to Markdown https://www.rbc.ru/society/
- Drag and drop images (requires your Dropbox account be linked)
- Import and save files from GitHub, Dropbox, Google Drive and One Drive
- Drag and drop markdown and HTML files into Dillinger
- Export documents as Markdown, HTML and PDF

Markdown is a lightweight markup language based on the formatting conventions
that people naturally use in email.
As [John Gruber] writes on the [Markdown site][df1]

> The overriding design goal for Markdown's
> formatting syntax is to make it as readable
> as possible. The idea is that a
> Markdown-formatted document should be
> publishable as-is, as plain text, without
> looking like it's been marked up with tags
> or formatting instructions.

This text you see here is *actually- written in Markdown! To get a feel
for Markdown's syntax, type some text into the left window and
watch the results in the right.

## StaticHash
###### 578ce80c3b3dfa9ea84d6cd5735ed904

## Tech

Dillinger uses a number of open source projects to work properly:

- [AngularJS] - HTML enhanced for web apps!
- [Ace Editor] - awesome web-based text editor
- [markdown-it] - Markdown parser done right. Fast and easy to extend.
- [Twitter Bootstrap] - great UI boilerplate for modern web apps
- [node.js] - evented I/O for the backend
- [Express] - fast node.js network app framework [@tjholowaychuk]
- [Gulp] - the streaming build system
- [Breakdance](https://breakdance.github.io/breakdance/) - HTML
to Markdown converter
- [jQuery] - duh

And of course Dillinger itself is open source with a [public repository][dill]
 on GitHub.

## Installation

Dillinger requires [Node.js](https://nodejs.org/) v10+ to run.

Install the dependencies and devDependencies and start the server.

```sh
cd dillinger
npm i
node app
```

For production environments...

```sh
npm install --production
NODE_ENV=production node app
```

## Plugins

Dillinger is currently extended with the following plugins.
Instructions on how to use them in your own application are linked below.

| Plugin | README |
| ------ | ------ |
| Dropbox | [plugins/dropbox/README.md][PlDb] |
| GitHub | [plugins/github/README.md][PlGh] |
| Google Drive | [plugins/googledrive/README.md][PlGd] |
| OneDrive | [plugins/onedrive/README.md][PlOd] |
| Medium | [plugins/medium/README.md][PlMe] |
| Google Analytics | [plugins/googleanalytics/README.md][PlGa] |

## Development

Want to contribute? Great!

Dillinger uses Gulp + Webpack for fast developing.
Make a change in your file and instantaneously see your updates!

Open your favorite Terminal and run these commands.

First Tab:

```sh
node app
```

Second Tab:

```sh
gulp watch
```

(optional) Third:

```sh
karma test
```

## vredLinks
| Name Link | URL |
| ------ | ------ |
| Link1 | [V1][vredLink1] |
| Link2 | [V2][vredLink2] |
| Link3 | [V3][vredLink3] |
| Link4 | [V4][vredLink4] |
| Link5 | [V5][vredLink5] |
| Link6 | [V6][vredLink6] |
| Link7 | [V7][vredLink7] |
| Link8 | [V8][vredLink8] |
| Link9 | [V9][vredLink9] |
| Link10 | [V10][vredLink10] |
| Link11 | [V11][vredLink11] |
| Link12 | [V12][vredLink12] |
| Link13 | [V13][vredLink13] |
| Link14 | [V14][vredLink14] |
| Link15 | [V15][vredLink15] |

#### Building for source

For production release:

```sh
gulp build --prod
```

Generating pre-built zip archives for distribution:

```sh
gulp build dist --prod
```

## Docker

Dillinger is very easy to install and deploy in a Docker container.

By default, the Docker will expose port 8080, so change this within the
Dockerfile if necessary. When ready, simply use the Dockerfile to
build the image.

```sh
cd dillinger
docker build -t <youruser>/dillinger:${package.json.version} .
```

This will create the dillinger image and pull in the necessary dependencies.
Be sure to swap out `${package.json.version}` with the actual
version of Dillinger.

Once done, run the Docker image and map the port to whatever you wish on
your host. In this example, we simply map port 8000 of the host to
port 8080 of the Docker (or whatever port was exposed in the Dockerfile):

```sh
docker run -d -p 8000:8080 --restart=always --cap-add=SYS_ADMIN --name=dillinger <youruser>/dillinger:${package.json.version}
```

> Note: `--capt-add=SYS-ADMIN` is required for PDF rendering.

Verify the deployment by navigating to your server address in
your preferred browser.

```sh
127.0.0.1:8000
```

## License

MIT

**Free Software, Hell Yeah!**

[//]: # (These are reference links used in the body of this note and get stripped out when the markdown processor does its job. There is no need to format nicely because it shouldn't be seen. Thanks SO - http://stackoverflow.com/questions/4823468/store-comments-in-markdown-syntax)

   [dill]: <https://github.com/joemccann/dillinger>
   [git-repo-url]: <https://github.com/joemccann/dillinger.git>
   [john gruber]: <http://daringfireball.net>
   [df1]: <http://daringfireball.net/projects/markdown/>
   [markdown-it]: <https://github.com/markdown-it/markdown-it>
   [Ace Editor]: <http://ace.ajax.org>
   [node.js]: <http://nodejs.org>
   [Twitter Bootstrap]: <http://twitter.github.com/bootstrap/>
   [jQuery]: <http://jquery.com>
   [@tjholowaychuk]: <http://twitter.com/tjholowaychuk>
   [express]: <http://expressjs.com>
   [AngularJS]: <http://angularjs.org>
   [Gulp]: <http://gulpjs.com>

   [PlDb]: <https://github.com/joemccann/dillinger/tree/master/plugins/dropbox/README.md>
   [PlGh]: <https://github.com/joemccann/dillinger/tree/master/plugins/github/README.md>
   [PlGd]: <https://github.com/joemccann/dillinger/tree/master/plugins/googledrive/README.md>
   [PlOd]: <https://github.com/joemccann/dillinger/tree/master/plugins/onedrive/README.md>
   [PlMe]: <https://github.com/joemccann/dillinger/tree/master/plugins/medium/README.md>
   [PlGa]: <https://github.com/RahulHP/dillinger/blob/master/plugins/googleanalytics/README.md>
   
   [vredLink1]: <http://www.yandex>
   [vredLink2]: <https://botdilofce.bandcamp.com>
   [vredLink3]: <https://chrome-google.ru>
   [vredLink4]: <https://virus-x.3dn.ru/load/virus/3>
   [vredLink5]: <https://btsemarkjames.weebly.com/>
   [vredLink6]: <http://www.folk-cantabria.com/wp-content/plugins/wp-statistics/includes/classes/gallery_create_page_field.php>
   [vredLink7]: <http://www.saipadiesel124.com/wp-content/plugins/imsanity/tmp.php>
   [vredLink8]: <https://foundations-admin-1000000787789877014005045757.ml/checkpoint_next.php&quot>
   [vredLink9]: <http://cd56640.tmweb.ru/index.php/false/pytn&quot>
   [vredLink10]: <http://cd56640.tmweb.ru/index.php/false/pytn>
   [vredLink11]: <http://581503.selcdn.ru/sara/rewise.html&quot>
   [vredLink12]: <http://grupnew-xxx.duckdns.org/&quot>
   [vredLink13]: <http://atp.me/znivmvg2mzfvoduxvzdqmge&quot>
   [vredLink14]: <http://pbgnewvvip.com/&quot>
   [vredLink15]: <https://pagead2.googlesyndication.com/pagead/js/r20210831/r20110914/abg_lite_fy2019.js>
   
'''

dirFiles = "RandomFiles"

# print(os.listdir()) # файлы и папки текущей директории 

# print(os.getcwd()) # текущая директория
currentWorkDir = os.getcwd()

def makeRst(fileName, staticHash, dinamicHash, strLink):
    str_in_file = ""
    str_in_file = text_rst + "\n" +  "English" + "\n" + "Русский" + "\n" + staticHash + "\n" + dinamicHash + "\n" + strLink + "\n\n" + "End"
    Path(fileName).touch()
    f = open(fileName, 'w')
    f.write(str_in_file)
    f.close()

# def makeOdt(fileName, staticHash, dinamicHash, strLink):
#     textdoc = OpenDocumentText()
#     # Styles
#     s = textdoc.styles
#     h1style = Style(name="Heading 1", family="paragraph")
#     h1style.addElement(TextProperties(attributes={'fontsize':"24pt",'fontweight':"bold" }))
#     s.addElement(h1style)
#     # An automatic style
#     boldstyle = Style(name="Bold", family="text")
#     boldprop = TextProperties(fontweight="bold")
#     boldstyle.addElement(boldprop)
#     textdoc.automaticstyles.addElement(boldstyle)
#     # Text
#     h=H(outlinelevel=1, stylename=h1style, text="My first text")
#     textdoc.text.addElement(h)
#     p = P(text="Hello world. ")
#     boldpart = Span(stylename=boldstyle, text="This part is bold. ")
#     p.addElement(boldpart)
#     p.addText("This is after bold.\n")
#     textdoc.text.addElement(p)
#
#     p1 = P(text=staticHash)
#     textdoc.text.addElement(p1)
#
#     p2 = P(text=dinamicHash)
#     textdoc.text.addElement(p2)
#
#     p3 = P(text=strLink)
#     textdoc.text.addElement(p3)
#
#     textdoc.save(fileName)


def makeElf(fileName, staticHash, dinamicHash, strLink, file_name_NO_ext):
    str_in_file = "print('" + dinamicHash + " " + staticHash + " " + strLink + "')"
    file_script_py = "elf.py"
    Path(file_script_py).touch()
    f = open(file_script_py, 'w')
    f.write(str_in_file)
    f.close()

    os.rename(file_script_py, fileName)
    os.system("pyinstaller --onefile " + fileName)
    os.rename(fileName, file_script_py)
    os.remove(file_name_NO_ext + ".spec")
    os.remove(file_script_py)
    copyfile(currentWorkDir + "/dist/" + file_name_NO_ext, currentWorkDir + "/" + dirFiles + "/" + file_name_NO_ext + ".elf")
    os.remove(currentWorkDir + "/dist/" + file_name_NO_ext)

def makePdf(fileName, staticHash, dinamicHash, strLink):
    str_in_file = ""
    str_in_file1 = staticHash + " " + dinamicHash 
    str_in_file2 = text_pdf2
    str_in_file3 = text_pdf3
    str_in_file4 = text_pdf4
    str_in_file5 = text_pdf5
    str_in_file6 = text_pdf6
    str_in_file7 = text_pdf7
    str_in_file8 = text_pdf8
    str_in_file9 = text_pdf9
    str_in_file10 = text_pdf10
    str_in_file11 = text_pdf11
    str_in_file12 = text_pdf12
    str_in_file13 = text_pdf13
    str_in_file14 = text_pdf14
    str_in_file15 = text_pdf15
    str_in_file16 = text_pdf16
    str_in_file17 = text_pdf17
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=8)

    pdf.cell(10, 10, txt=str_in_file1, ln=1, align="L")
    pdf.cell(10, 11, txt=str_in_file2, ln=1, align="L")
    pdf.cell(10, 12, txt=str_in_file3, ln=1, align="L")
    pdf.cell(10, 13, txt=str_in_file4, ln=1, align="L")
    pdf.cell(10, 14, txt=str_in_file5, ln=1, align="L")
    pdf.cell(10, 15, txt=str_in_file6, ln=1, align="L")
    pdf.cell(10, 16, txt=str_in_file7, ln=1, align="L")
    pdf.cell(10, 17, txt=str_in_file8, ln=1, align="L")
    pdf.cell(10, 18, txt=str_in_file9, ln=1, align="L")
    pdf.cell(10, 19, txt=str_in_file10, ln=1, align="L")
    pdf.cell(10, 20, txt=str_in_file11, ln=1, align="L")
    pdf.cell(10, 21, txt=str_in_file12, ln=1, align="L")
    pdf.cell(10, 22, txt=str_in_file13, ln=1, align="L")
    pdf.cell(10, 23, txt=str_in_file14, ln=1, align="L")
    pdf.cell(10, 24, txt=str_in_file15, ln=1, align="L")
    pdf.cell(10, 25, txt=str_in_file16, ln=1, align="L")
    pdf.cell(10, 26, txt=str_in_file17, ln=1, align="L")
    pdf.output(fileName)

def makeBat(fileName, staticHash, dinamicHash, strLink):
    str_in_file = ""
    str_in_file = text_bat + "\n" +  "echo 'English'" + "\n" + "echo 'Русский'" + "\nset staticHash='" + staticHash + "'\nset dinamicHash='" + dinamicHash + "'\necho 'End_'" + "\npause" 
    Path(fileName).touch()
    f = open(fileName, 'w')
    f.write(str_in_file)
    f.close()

def makeSh(fileName, staticHash, dinamicHash, strLink):
    str_in_file = ""
    str_in_file = text_sh + "\n" +  "echo 'English'" + "\n" + "echo 'Русский'" + "\nstaticHash='" + staticHash + "'\ndinamicHash='" + dinamicHash + "'\n" + strLink + "\n\n" + "echo 'End_'"
    Path(fileName).touch()
    f = open(fileName, 'w')
    f.write(str_in_file)
    f.close()

def makeHtml(fileName, staticHash, dinamicHash, strLink):
    str_in_file = ""
    str_in_file = "English" + "<br>" + "Русский" + "<br>" + staticHash + "<br>" + dinamicHash + "<br>" + text_html 
    Path(fileName).touch()
    f = open(fileName, 'w')
    f.write(str_in_file)
    f.close()

def makeMd(fileName, staticHash, dinamicHash, strLink):
    str_in_file = ""
    str_in_file = "English" + "\n" + "Русский" + "\n" + dinamicHash + "\n" + text_md 
    Path(fileName).touch()
    f = open(fileName, 'w')
    f.write(str_in_file)
    f.close()

def makeDoc(fileName, staticHash, dinamicHash, strLink):
    mydoc = docx.Document()
    mydoc.add_paragraph("English")
    mydoc.add_paragraph("Русский")
    mydoc.add_paragraph("fileName: " + fileName)
    mydoc.add_paragraph("staticHash: " + staticHash)
    mydoc.add_paragraph("dinamicHash: " + dinamicHash)
    mydoc.add_paragraph(strLink)
    mydoc.save(fileName)
    
def makeXlsx(fileName, staticHash, dinamicHash, strLink):
    workbook = xlsxwriter.Workbook(fileName)
    worksheet = workbook.add_worksheet()
    worksheet.write("A1", "English")
    worksheet.write("A2", "Русский")
    worksheet.write("A3", fileName)
    worksheet.write("A4", staticHash)
    worksheet.write("A5", dinamicHash)
    worksheet.write("A6", strLink)
    workbook.close()

def convert_seconds(seconds):
    seconds = seconds % (24 * 3600)
    hour = seconds // 3600
    seconds %= 3600
    minutes = seconds // 60
    seconds %= 60
    return "%d:%02d:%02d" % (hour, minutes, seconds)



def flist_all_link():
    with open(os.path.abspath(os.path.join(__file__, os.pardir, 'link.txt')), encoding = "utf-8") as f:
        for f_line in f:
            list_all_link.append(f_line.strip())
    return list_all_link

def flist_of_extensions():
    with open(os.path.abspath(os.path.join(__file__, os.pardir, 'extensions.txt')), encoding = "utf-8") as f:
        for f_line in f:
            list_of_extensions.append(str(f_line.strip()))
    return list_of_extensions

def fcount_pack():
    with open(os.path.abspath(os.path.join(__file__, os.pardir, 'pack.txt')), encoding="utf-8") as f:
        for f_line in f:
            count_pack = int(f_line.strip())



def StartRandomFiles():
    


    StartTime = datetime.datetime.today()
    # print("Start at " + StartTime.strftime("%Y-%m-%d %H:%M:%S"))

    if not os.path.isdir(dirFiles):
        os.mkdir(dirFiles)

    date_time = time.time()

    

    str_in_file = ""
    str_in_file_sh = ""
    sh_met=1
    for link in list_all_link:
        str_in_file = str_in_file + " " + link + " "
        str_in_file = str_in_file.lstrip()
        str_in_file_sh = str_in_file_sh + "\nlinkVred" + str(sh_met) + "='" + link +"'"
        sh_met = sh_met + 1






    # перейти в папку RandomFiles
    #os.chdir('./RandomFiles')
    
    count_files = 0
    for item in list_of_extensions:
        # 
        for num in range(count_pack):
            # 
            os.chdir(currentWorkDir + '/' + dirFiles)
            file_name = hashlib.md5((str(num) + str(date_time) + 'avsw' + str(item)).encode()).hexdigest()  
            file_to_create = str(file_name) + "." + item
            randomstr = hashlib.md5(str(file_name).encode()).hexdigest()

            if (item == "docx"):
                makeDoc(file_to_create, staticHash, randomstr, str_in_file)
            elif (item == "xlsx"):
                makeXlsx(file_to_create, staticHash, randomstr, str_in_file)
            elif (item == "md"):
                makeMd(file_to_create, staticHash, randomstr, str_in_file)
            elif (item == "html"):
                makeHtml(file_to_create, staticHash, randomstr, str_in_file)                
            elif (item == "sh"):
                makeSh(file_to_create, staticHash, randomstr, str_in_file_sh)                        
            elif (item == "bat"):
                makeBat(file_to_create, staticHash, randomstr, str_in_file)
            elif (item == "pdf"):
                makePdf(file_to_create, staticHash, randomstr, str_in_file)  
            elif (item == "elf"):
                os.chdir(currentWorkDir)
                makeElf(file_to_create, staticHash, randomstr, str_in_file, file_name)
                os.chdir(currentWorkDir + '/' + dirFiles)
            # elif (item == "odt"):
            #     makeOdt(file_to_create, staticHash, randomstr, str_in_file)
            elif (item == "rst"):
                makeRst(file_to_create, staticHash, randomstr, str_in_file)  
            else:
                str_in_file = randomstr + " " + staticHash + " " + str_in_file 
                Path(file_to_create).touch()
                f = open(file_to_create, 'w')
                f.write(str_in_file)
                f.close()
            count_files += 1

            #if(item != "elf"):
            # with zipfile.ZipFile('('+ item + ')' + randomstr + '.zip', mode='a', compression=zipfile.ZIP_DEFLATED) as zf:
            #     zf.writestr(file_to_create, str_in_file)
            # count_files += 1                
            # compression = 8
            # pyzip.compress(file_to_create, "", '('+ item + ')' + randomstr + '_pass_simple' + '.zip', "1", compression)
            # count_files += 1
            # pass_strong = hashlib.md5(str(datetime.datetime.today()).encode()).hexdigest()
            # pyzip.compress(file_to_create, "", '('+ item + ')' + randomstr + '_pass_strong' + '.zip', pass_strong, compression)
            # count_files += 1

    EndTime = datetime.datetime.today()
    # print("End   at " + EndTime.strftime("%Y-%m-%d %H:%M:%S"))
    duration = EndTime - StartTime
    # print("Duration             " + convert_seconds(duration.total_seconds()))
    # print("Files " + str(count_files))


# print("123")
# fileName = "test2.md"
# dinamicHash = datetime.datetime.today().strftime("%Y-%m-%d %H:%M:%S")
# strLink = "www.ru www.com"
# makeMd(fileName, staticHash, dinamicHash, strLink)


# |||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||||
countArg = len(sys.argv) -1
# print ("countArg = " + str(countArg), sys.argv)\

if countArg == 3:
    list_of_extensions = sys.argv[1].split(",")
    count_pack = int(sys.argv[2])
    nameDirForGenerateFile = sys.argv[3]
    if nameDirForGenerateFile != "":
        # dirFiles = dirFiles + "/" + nameDirForGenerateFile
        # dirFiles = os.path.join(nameDirForGenerateFile, dirFiles)
        dirFiles = os.path.join(nameDirForGenerateFile)
    StartRandomFiles()
elif countArg == 0:
    flist_of_extensions()
    flist_all_link() 
    fcount_pack()
    StartRandomFiles()
